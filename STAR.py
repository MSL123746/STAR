from __future__ import annotations

from io import BytesIO

import streamlit as st
from docx import Document


st.set_page_config(
    page_title="AI-STAR Interview Prep",
    page_icon="*",
    layout="wide",
    initial_sidebar_state="collapsed",
)


APP_CSS = """
<style>
:root {
    --page: #f3f7fa;
    --title: #20344b;
    --help: #495869;
    --field-bg: #232d39;
    --field-border: #17202a;
    --field-text: #ffffff;
    --panel-bg: #e7f0f8;
    --panel-border: #d4dee8;
    --generate-bg: #d7dde3;
    --generate-border: #bfc8d0;
    --generate-text: #70757c;
    --word-bg: #2b6f4d;
    --word-hover: #245e41;
}

html, body, [class*="css"] {
    background: var(--page);
    color: #111827;
    font-family: Arial, Helvetica, sans-serif;
}

.block-container {
    padding: 0.5in 0.35rem 0.6rem 0.35rem;
    max-width: 100% !important;
}

.header {
    display: flex;
    align-items: center;
    gap: 0.5rem;
    font-size: 2.9rem;
    font-weight: 700;
    color: var(--title);
    margin-bottom: 0.65rem;
}

.header-icon {
    font-size: 2.3rem;
}

.block-label {
    font-size: 1.84rem;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0.02em;
    margin: 0.3rem 0 0.14rem 0;
}

.help-text {
    font-size: 1.6rem;
    line-height: 1.18;
    color: var(--help);
    margin-bottom: 0.08rem;
}

.right-title {
    font-size: 1.9rem;
    font-weight: 800;
    text-transform: uppercase;
    line-height: 1.15;
    margin-bottom: 0.3rem;
}

.right-help {
    font-size: 1.6rem;
    line-height: 1.18;
    color: #334155;
    margin-bottom: 0.28rem;
}

.right-box {
    background: var(--panel-bg);
    border: 1px solid var(--panel-border);
    border-radius: 6px;
    min-height: 780px;
    padding: 0.85rem;
    white-space: pre-wrap;
    color: #1f2937;
    font-size: 1.6rem;
    line-height: 1.2;
}

.stTextArea textarea {
    background: var(--field-bg) !important;
    color: #000000 !important;
    border: 1px solid var(--field-border) !important;
    border-radius: 4px !important;
    box-shadow: none !important;
    font-size: 1.64rem !important;
    line-height: 1.15 !important;
}

.stTextArea textarea::placeholder {
    color: rgba(255, 255, 255, 0.42) !important;
}

.stTextArea label,
div[data-testid="stCheckbox"] label > div:first-child + div label,
[data-testid="stTextInput"] label {
    display: none !important;
}

div[data-testid="stCheckbox"] label p,
div[data-testid="stCheckbox"] label span {
    font-size: 1.8rem !important;
    line-height: 1.15 !important;
    color: #334155 !important;
}

.stButton > button {
    width: auto !important;
    border-radius: 6px !important;
    border: 1px solid var(--generate-border) !important;
    background: var(--generate-bg) !important;
    color: var(--generate-text) !important;
    font-weight: 700 !important;
    font-size: 2.2rem !important;
    padding: 1rem 1.2rem !important;
}

.stButton > button span,
.stButton > button p,
.stButton > button div {
    font-size: 2.2rem !important;
    line-height: 1.1 !important;
}

.stButton > button:hover {
    background: #cfd5dc !important;
}

.stDownloadButton > button {
    width: auto !important;
    border-radius: 6px !important;
    border: 1px solid transparent !important;
    background: var(--word-bg) !important;
    color: white !important;
    font-weight: 700 !important;
    font-size: 2.2rem !important;
    padding: 1rem 1.2rem !important;
}

.stDownloadButton > button span,
.stDownloadButton > button p,
.stDownloadButton > button div {
    font-size: 2.2rem !important;
    line-height: 1.1 !important;
}

.stDownloadButton > button:hover {
    background: var(--word-hover) !important;
}

.footer-note {
    color: #7b8794;
    font-size: 1.36rem;
    margin-top: 0.5rem;
}

.action-spacer {
    height: 2.3rem;
}

div[data-testid="column"] .stButton,
div[data-testid="column"] .stDownloadButton {
    display: flex;
    justify-content: flex-start;
}
</style>
"""


def build_followups(question: str, situation: str, task: str, action: str, result: str) -> list[str]:
    if not any([question.strip(), situation.strip(), task.strip(), action.strip(), result.strip()]):
        return []

    prompts = [
        "What follow-up questions would an interviewer ask about your decision making?",
        "What measurable outcome best proves the impact of your example?",
    ]
    if situation.strip():
        prompts[0] = f"What made this situation challenging: {situation.strip()[:85]}?"
    if result.strip():
        prompts[1] = f"How would you quantify this result more clearly: {result.strip()[:85]}?"
    return prompts


def build_summary(question: str, situation: str, task: str, action: str, result: str) -> str:
    lines: list[str] = []
    if question.strip():
        lines.append(f"Behavioral question: {question.strip()}")
    if situation.strip():
        lines.append(f"Situation: {situation.strip()}")
    if task.strip():
        lines.append(f"Task: {task.strip()}")
    if action.strip():
        lines.append(f"Action: {action.strip()}")
    if result.strip():
        lines.append(f"Result: {result.strip()}")
    return "\n".join(lines)


def create_docx_bytes(answer_text: str) -> bytes:
    document = Document()
    document.add_heading("AI-STAR Interview Prep Answer", level=1)
    for paragraph in answer_text.splitlines() or [""]:
        document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


st.markdown(APP_CSS, unsafe_allow_html=True)
st.markdown(
    '<div class="header"><span class="header-icon">⭐</span><span>AI-STAR Interview Prep | Behavioral Question &amp; Answer Builder</span></div>',
    unsafe_allow_html=True,
)

left_col, right_col = st.columns([1.72, 1], gap="small")

with left_col:
    st.markdown('<div class="block-label">Enter Behavioral Question</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="help-text">Enter the behavioral-based question you want to answer (e.g., "Tell me about a time you faced a difficult challenge and how you handled it").</div>',
        unsafe_allow_html=True,
    )
    question = st.text_area("Behavioral Question", height=74, placeholder="", label_visibility="collapsed")

    st.markdown('<div class="block-label">Craft Your STAR Narrative</div>', unsafe_allow_html=True)

    st.markdown('<div class="help-text"><strong>(S) Situation (Background &amp; Context)</strong><br/>Describe the specific situation where the challenge occurred. (Where, When, Who was involved?)</div>', unsafe_allow_html=True)
    situation = st.text_area("Situation", height=82, placeholder="", label_visibility="collapsed")

    st.markdown('<div class="help-text"><strong>(T) Task (The Responsibility &amp; Goal)</strong><br/>Define the specific challenge or task you needed to address. What was your main objective?</div>', unsafe_allow_html=True)
    task = st.text_area("Task", height=82, placeholder="", label_visibility="collapsed")

    st.markdown('<div class="help-text"><strong>(A) Action (Detailed Steps taken)</strong><br/>Outline the specific steps you took to tackle the challenge. (Your technical maneuvers and leadership actions.)</div>', unsafe_allow_html=True)
    action = st.text_area("Action", height=96, placeholder="", label_visibility="collapsed")

    st.markdown('<div class="help-text"><strong>(R) Result (The Outcome &amp; Impact)</strong><br/>Describe the final outcome and the value delivered. (Quantify the success and lasting stability.)</div>', unsafe_allow_html=True)
    result = st.text_area("Result", height=96, placeholder="", label_visibility="collapsed")

with right_col:
    st.markdown('<div class="right-title">Your STAR Narrative</div>', unsafe_allow_html=True)

    if "answer_text" not in st.session_state:
        st.session_state.answer_text = ""

    rendered_answer = st.session_state.answer_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    st.markdown(f'<div class="right-box">{rendered_answer}</div>', unsafe_allow_html=True)

left_actions, right_actions = st.columns([1.72, 1], gap="small")

with left_actions:
    st.markdown('<div class="block-label">AI Enhancement</div>', unsafe_allow_html=True)
    enable_followups = st.checkbox("Get 2 customized AI follow-up questions based on my STAR narrative.", value=True)
    st.markdown('<div class="action-spacer"></div>', unsafe_allow_html=True)
    generate_clicked = st.button("Generate Final Answer with AI")

with right_actions:
    st.markdown('<div class="action-spacer"></div>', unsafe_allow_html=True)
    docx_bytes = create_docx_bytes(st.session_state.answer_text)
    st.download_button(
        "Download Answer as MS Word Doc (.docx)",
        data=docx_bytes,
        file_name="ai_star_answer.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download-docx",
    )

if generate_clicked:
    followups = build_followups(question, situation, task, action, result) if enable_followups else []
    summary = build_summary(question, situation, task, action, result)
    parts: list[str] = []
    if followups:
        for item in followups:
            parts.append(item)
    if summary:
        if parts:
            parts.append("")
        parts.append(summary)
    st.session_state.answer_text = "\n".join(parts)
    st.rerun()

st.markdown('<div class="footer-note">Tip: Fill in the STAR sections, click generate, and then download the answer as a Word document.</div>', unsafe_allow_html=True)

